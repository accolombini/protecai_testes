# -*- coding: utf-8 -*-
"""
Gerador de documentação DOCX explicativa para códigos normalizados.

Este script:
1. Lê os dados normalizados (do Excel gerado ou diretamente dos arquivos)
2. Agrupa valores únicos e suas interpretações
3. Gera documento DOCX com tabelas explicativas
4. Organiza por tipo de código e nível de confiança

Estrutura do documento:
- Sumário executivo
- Tabela de códigos ANSI encontrados
- Tabela de tipos de proteção
- Tabela de model numbers
- Tabela de valores não interpretados (para revisão)
- Estatísticas do processamento

    ||> documentacao_codigos_latest.docx

Autor: Sistema ProtecAI
Data: 2025-10-03
"""

from __future__ import annotations
import sys
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple
import pandas as pd
from datetime import datetime
from collections import defaultdict

# Imports para DOCX
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

# Importar o gerador de Excel (para reutilizar dados)
try:
    from .generate_normalized_excel import NormalizedExcelGenerator
    from .code_parser import TokenType
except ImportError:
    # Para execução direta
    sys.path.append(str(Path(__file__).parent))
    from generate_normalized_excel import NormalizedExcelGenerator
    from code_parser import TokenType


# Configurações de diretórios
THIS_FILE = Path(__file__).resolve()
REPO_ROOT = THIS_FILE.parents[2]
INPUT_DIR = REPO_ROOT / "outputs" / "atrib_limpos"
OUTPUT_DIR = REPO_ROOT / "outputs" / "doc"

# Garantir que o diretório de saída existe
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)


class DocxDocumentationGenerator:
    """Gerador de documentação DOCX para códigos normalizados."""
    
    def __init__(self):
        self.excel_generator = NormalizedExcelGenerator()
        self.distinct_values: Optional[pd.DataFrame] = None
        self.stats: Dict[str, Any] = {}
    
    def process_data(self, input_dir: Path = INPUT_DIR) -> None:
        """
        Processa os dados e prepara para geração da documentação.
        
        Args:
            input_dir: Diretório com arquivos a processar
        """
        print("[INFO] Processando dados para documentação...")
        
        # Processar usando o gerador de Excel
        self.excel_generator.process_directory(input_dir)
        
        if not self.excel_generator.normalized_rows:
            raise ValueError("Nenhum dado foi processado.")
        
        # Obter valores distintos
        self.distinct_values = self.excel_generator.get_distinct_values_for_doc()
        self.stats = self.excel_generator.stats.copy()
        
        print(f"[INFO] Processamento concluído. {len(self.distinct_values)} valores únicos encontrados.")
    
    def generate_documentation(self, output_path: Optional[Path] = None) -> Path:
        """
        Gera a documentação DOCX completa.
        
        Args:
            output_path: Caminho de saída (opcional)
            
        Returns:
            Caminho do arquivo gerado
        """
        if self.distinct_values is None:
            raise ValueError("Execute process_data() primeiro.")
        
        if output_path is None:
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            output_path = OUTPUT_DIR / f"documentacao_codigos_{timestamp}.docx"
        
        # Criar documento
        doc = Document()
        self._setup_document_style(doc)
        
        # Adicionar conteúdo
        self._add_title_page(doc)
        self._add_executive_summary(doc)
        self._add_ansi_codes_section(doc)
        self._add_protection_types_section(doc)
        self._add_model_numbers_section(doc)
        self._add_uninterpreted_values_section(doc)
        self._add_statistics_section(doc)
        
        # Salvar documento
        doc.save(output_path)
        
        print(f"[OK] Documentação DOCX gerada: {output_path}")
        return output_path
    
    def _setup_document_style(self, doc: Document) -> None:
        """Configura estilos do documento."""
        # Configurar fonte padrão
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
        
        # Configurar estilo de cabeçalho
        heading1_style = doc.styles['Heading 1']
        heading1_style.font.name = 'Calibri'
        heading1_style.font.size = Pt(16)
        heading1_style.font.color.rgb = RGBColor(0, 51, 102)  # Azul escuro
    
    def _add_title_page(self, doc: Document) -> None:
        """Adiciona página de título."""
        title = doc.add_heading('Documentação de Códigos de Proteção', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = doc.add_heading('Análise e Normalização de Códigos Multivalorados', level=2)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Espaço
        
        info_para = doc.add_paragraph()
        info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info_para.add_run("Sistema ProtecAI\\n")
        info_para.add_run(f"Processado em: {datetime.now().strftime('%d/%m/%Y às %H:%M')}\\n")
        info_para.add_run(f"Base: outputs/atrib_limpos/")
        
        doc.add_page_break()
    
    def _add_executive_summary(self, doc: Document) -> None:
        """Adiciona sumário executivo."""
        doc.add_heading('Sumário Executivo', level=1)
        
        doc.add_paragraph(
            "Este documento apresenta a análise e normalização automática de códigos "
            "multivalorados encontrados nos dados de configuração de relés de proteção. "
            "O sistema identificou e segmentou códigos como '52-MP-20', model numbers "
            "como 'P241311B2M0600J' e outras estruturas codificadas em suas partes atômicas."
        )
        
        # Tabela de resumo
        summary_table = doc.add_table(rows=1, cols=2)
        summary_table.style = 'Light List Accent 1'
        summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Cabeçalho
        hdr_cells = summary_table.rows[0].cells
        hdr_cells[0].text = 'Métrica'
        hdr_cells[1].text = 'Valor'
        
        # Dados do resumo
        summary_data = [
            ('Arquivos processados', str(self.stats.get('files_processed', 0))),
            ('Total de valores analisados', str(self.stats.get('total_values', 0))),
            ('Valores multivalorados', str(self.stats.get('multivalued_found', 0))),
            ('Valores atômicos', str(self.stats.get('atomic_values', 0))),
            ('Tokens extraídos', str(self.stats.get('tokens_extracted', 0))),
            ('Valores únicos identificados', str(len(self.distinct_values)))
        ]
        
        for metric, value in summary_data:
            row_cells = summary_table.add_row().cells
            row_cells[0].text = metric
            row_cells[1].text = value
    
    def _add_ansi_codes_section(self, doc: Document) -> None:
        """Adiciona seção de códigos ANSI."""
        doc.add_heading('Códigos ANSI/IEEE Identificados', level=1)
        
        doc.add_paragraph(
            "Esta seção lista todos os códigos ANSI/IEEE encontrados nos dados. "
            "Estes códigos seguem o padrão internacional para funções de proteção."
        )
        
        # Filtrar códigos ANSI
        ansi_data = self._get_values_by_pattern_or_content('ansi')
        
        if ansi_data.empty:
            doc.add_paragraph("Nenhum código ANSI foi identificado nos dados processados.")
            return
        
        # Criar tabela
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Light List Accent 1'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Cabeçalho
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Código'
        hdr_cells[1].text = 'Significado'
        hdr_cells[2].text = 'Valor Original'
        hdr_cells[3].text = 'Confiança'
        
        # Adicionar dados
        for _, row in ansi_data.iterrows():
            row_cells = table.add_row().cells
            row_cells[0].text = self._extract_ansi_from_interpretation(row['interpretacao'])
            row_cells[1].text = self._extract_ansi_meaning(row['interpretacao'])
            row_cells[2].text = str(row['valor_original'])
            row_cells[3].text = f"{row['confianca_geral']:.1%}"
    
    def _add_protection_types_section(self, doc: Document) -> None:
        """Adiciona seção de tipos de proteção."""
        doc.add_heading('Tipos de Proteção Identificados', level=1)
        
        doc.add_paragraph(
            "Esta seção apresenta os tipos de proteção (MP, EF, OC, etc.) "
            "identificados nos códigos estruturados."
        )
        
        # Filtrar tipos de proteção
        protection_data = self._get_values_by_pattern_or_content('protection')
        
        if protection_data.empty:
            doc.add_paragraph("Nenhum tipo de proteção específico foi identificado.")
            return
        
        # Criar tabela
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Light List Accent 1'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Cabeçalho
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Código'
        hdr_cells[1].text = 'Significado'
        hdr_cells[2].text = 'Valor Original'
        hdr_cells[3].text = 'Confiança'
        
        # Adicionar dados
        for _, row in protection_data.iterrows():
            row_cells = table.add_row().cells
            prot_code, prot_meaning = self._extract_protection_info(row['interpretacao'])
            row_cells[0].text = prot_code
            row_cells[1].text = prot_meaning
            row_cells[2].text = str(row['valor_original'])
            row_cells[3].text = f"{row['confianca_geral']:.1%}"
    
    def _add_model_numbers_section(self, doc: Document) -> None:
        """Adiciona seção de model numbers."""
        doc.add_heading('Model Numbers e Códigos de Produto', level=1)
        
        doc.add_paragraph(
            "Esta seção documenta os model numbers e códigos de produto "
            "identificados e suas respectivas segmentações."
        )
        
        # Filtrar model numbers
        model_data = self._get_values_by_pattern_or_content('model')
        
        if model_data.empty:
            doc.add_paragraph("Nenhum model number foi identificado.")
            return
        
        # Criar tabela
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Light List Accent 1'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Cabeçalho
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Model Number'
        hdr_cells[1].text = 'Segmentação'
        hdr_cells[2].text = 'Padrão Detectado'
        hdr_cells[3].text = 'Confiança'
        
        # Adicionar dados
        for _, row in model_data.iterrows():
            row_cells = table.add_row().cells
            row_cells[0].text = str(row['valor_original'])
            row_cells[1].text = str(row['interpretacao'])
            row_cells[2].text = str(row['padrao_detectado'])
            row_cells[3].text = f"{row['confianca_geral']:.1%}"
    
    def _add_uninterpreted_values_section(self, doc: Document) -> None:
        """Adiciona seção de valores não interpretados."""
        doc.add_heading('Valores Não Interpretados - Revisão Manual', level=1)
        
        doc.add_paragraph(
            "Esta seção lista valores que não foram automaticamente interpretados "
            "e requerem revisão manual para definição de significado."
        )
        
        # Filtrar valores com baixa confiança ou não interpretados
        uninterpreted = self.distinct_values[
            (self.distinct_values['confianca_geral'] < 0.5) | 
            (self.distinct_values['interpretacao'].str.contains('Não interpretado|não reconhecido|revisar', case=False, na=False))
        ]
        
        if uninterpreted.empty:
            doc.add_paragraph("Excelente! Todos os valores foram interpretados com alta confiança.")
            return
        
        # Criar tabela
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Light List Accent 1'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Cabeçalho
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Valor Original'
        hdr_cells[1].text = 'Interpretação Tentativa'
        hdr_cells[2].text = 'Ação Requerida'
        
        # Adicionar dados
        for _, row in uninterpreted.iterrows():
            row_cells = table.add_row().cells
            row_cells[0].text = str(row['valor_original'])
            row_cells[1].text = str(row['interpretacao'])
            
            if row['confianca_geral'] < 0.3:
                action = "Revisão completa necessária"
            elif row['confianca_geral'] < 0.5:
                action = "Verificar interpretação"
            else:
                action = "Confirmar significado"
            row_cells[2].text = action
    
    def _add_statistics_section(self, doc: Document) -> None:
        """Adiciona seção de estatísticas detalhadas."""
        doc.add_heading('Estatísticas Detalhadas', level=1)
        
        doc.add_paragraph(
            "Esta seção apresenta estatísticas detalhadas sobre o processamento "
            "e a distribuição dos tipos de códigos identificados."
        )
        
        # Estatísticas por tipo de token
        if hasattr(self.excel_generator, 'normalized_rows') and self.excel_generator.normalized_rows:
            df_all = pd.DataFrame(self.excel_generator.normalized_rows)
            
            # Distribuição por tipo de token (nova estrutura horizontal)
            doc.add_heading('Distribuição por Tipo de Token', level=2)
            
            # Coletar todos os tipos de tokens das colunas horizontais
            token_types = []
            for i in range(1, self.excel_generator.max_tokens + 1):
                type_col = f'token_{i}_type'
                if type_col in df_all.columns:
                    types = df_all[type_col].dropna()
                    types = types[types != ""]
                    token_types.extend(types.tolist())
            
            if token_types:
                from collections import Counter
                type_counts = Counter(token_types)
                
                # Criar tabela
                table = doc.add_table(rows=1, cols=2)
                table.style = 'Light List Accent 1'
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                
                # Cabeçalho
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Tipo de Token'
                hdr_cells[1].text = 'Quantidade'
                
                # Adicionar dados
                for token_type, count in type_counts.most_common():
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(token_type).replace('_', ' ').title()
                    row_cells[1].text = str(count)
            else:
                doc.add_paragraph("Nenhum token foi extraído dos valores processados.")
        
        # Observações finais
        doc.add_heading('Observações e Recomendações', level=2)
        
        total_values = self.stats.get('total_values', 0)
        multivalued = self.stats.get('multivalued_found', 0)
        atomic = self.stats.get('atomic_values', 0)
        
        if total_values > 0:
            multivalued_pct = (multivalued / total_values) * 100
            atomic_pct = (atomic / total_values) * 100
            
            observations = [
                f"• {multivalued_pct:.1f}% dos valores analisados são multivalorados e foram segmentados.",
                f"• {atomic_pct:.1f}% dos valores são atômicos e não requerem normalização.",
                "• Valores com confiança < 50% devem ser revisados manualmente.",
                "• Novos padrões identificados podem ser adicionados ao sistema para melhorar a cobertura.",
                "• O sistema pode ser expandido conforme novos tipos de equipamentos forem processados."
            ]
            
            for obs in observations:
                doc.add_paragraph(obs)
    
    def _get_values_by_pattern_or_content(self, filter_type: str) -> pd.DataFrame:
        """
        Filtra valores por tipo/padrão.
        
        Args:
            filter_type: 'ansi', 'protection', 'model', etc.
            
        Returns:
            DataFrame filtrado
        """
        if self.distinct_values is None or self.distinct_values.empty:
            return pd.DataFrame()
        
        if filter_type == 'ansi':
            # Buscar por códigos ANSI na interpretação
            mask = self.distinct_values['interpretacao'].str.contains(
                r'ANSI|52|51|50|67|27|59|81|86|87|94|49|46|32|40|64|79|68|25', 
                case=False, na=False, regex=True
            )
        elif filter_type == 'protection':
            # Buscar por tipos de proteção
            mask = self.distinct_values['interpretacao'].str.contains(
                r'Protection|Motor|Earth Fault|Overcurrent|Protection|MP|EF|OC', 
                case=False, na=False, regex=True
            )
        elif filter_type == 'model':
            # Buscar por model numbers (contém padrões específicos)
            mask = (
                self.distinct_values['padrao_detectado'].str.contains('model', case=False, na=False) |
                self.distinct_values['valor_original'].str.match(r'^[A-Z]\d', na=False)
            )
        else:
            # Retornar vazio para tipos não reconhecidos
            mask = pd.Series([False] * len(self.distinct_values))
        
        return self.distinct_values[mask]
    
    def _extract_ansi_from_interpretation(self, interpretation: str) -> str:
        """Extrai código ANSI da interpretação."""
        import re
        match = re.search(r'(\d{2}N?)\s*=', str(interpretation))
        return match.group(1) if match else "N/A"
    
    def _extract_ansi_meaning(self, interpretation: str) -> str:
        """Extrai significado ANSI da interpretação."""
        import re
        match = re.search(r'\d{2}N?\s*=\s*([^|]+)', str(interpretation))
        if match:
            return match.group(1).strip()
        return "Significado não extraído"
    
    def _extract_protection_info(self, interpretation: str) -> Tuple[str, str]:
        """Extrai informações de proteção da interpretação."""
        import re
        # Buscar padrão como "MP = Motor Protection"
        match = re.search(r'([A-Z]{2,4})\s*=\s*([^|]+)', str(interpretation))
        if match:
            return match.group(1).strip(), match.group(2).strip()
        return "N/A", "Não extraído"


def main() -> int:
    """Função principal do script."""
    try:
        print("[INFO] Iniciando geração de documentação DOCX...")
        
        # Criar o gerador
        generator = DocxDocumentationGenerator()
        
        # Processar dados
        generator.process_data()
        
        # Gerar documentação
        output_path = generator.generate_documentation()
        
        print(f"[SUCESSO] Documentação gerada: {output_path}")
        
        return 0
        
    except Exception as e:
        print(f"[ERRO] Falha na geração: {e}")
        import traceback
        traceback.print_exc()
        return 1


if __name__ == "__main__":
    sys.exit(main())